import os
import tkinter as tk
import datetime
from openpyxl import load_workbook
from docx import Document


# 02 16 区模板
def outputdocument1(nomDoc: str, placeDuFichier: str):
    for i in range(1, inputpage.max_row):
        if cell[i][11].value == "当天":
            if len(str(datetime.datetime.now().month)) == 1:
                datetimes = str(datetime.datetime.now().day) + "/0" + str(datetime.datetime.now().month) + "/" + str(
                    datetime.datetime.now().year)
            else:
                datetimes = str(datetime.datetime.now().day) + "/" + str(datetime.datetime.now().month) + "/" + str(
                    datetime.datetime.now().year)
        else:
            datetimes = cell[i][11].value
        textlist = [str(cell[i][9].value), str(cell[i][3].value), str(cell[i][10].value), str(datetimes),
                    str(cell[i][3].value), str(datetimes)]

        doc = Document("授权书\\" + nomDoc)
        count = 0
        for p in doc.paragraphs:
            if '****' in p.text:
                inline = p.runs
                for j in range(len(inline)):
                    if '****' in inline[j].text:
                        text = inline[j].text.replace('****', textlist[count])
                        inline[j].text = text
                        count += 1
        doc.save(placeDuFichier + "\\" + str(cell[i][0].value) + "-Mandat-" + cell[i][3].value.upper() + ".docx")


def outputdocument2(nomDoc: str, placeDuFichier: str):
    for i in range(1, inputpage.max_row):
        if cell[i][11].value == "当天":
            if len(str(datetime.datetime.now().month)) == 1:
                datetimes = str(datetime.datetime.now().day) + "/0" + str(datetime.datetime.now().month) + "/" + str(
                    datetime.datetime.now().year)
            else:
                datetimes = str(datetime.datetime.now().day) + "/" + str(datetime.datetime.now().month) + "/" + str(
                    datetime.datetime.now().year)
        else:
            datetimes = cell[i][11].value
        textlist = [str(cell[i][9].value), str(cell[i][3].value), str(cell[i][10].value), str(datetimes),
                    str(cell[i][3].value), str(datetimes)]

        doc = Document("授权书\\" + nomDoc)
        count = 0
        for p in doc.paragraphs:
            if '****' in p.text:
                inline = p.runs
                for j in range(len(inline)):
                    if '****' in inline[j].text:
                        text = inline[j].text.replace('****', textlist[count])
                        inline[j].text = text
                        count += 1
                    
        doc.save(placeDuFichier + "\\" + cell[i][3].value.upper() + "-Mandat-" + str(cell[i][0].value) + ".docx")


# Chatillon 模板
def outputdocument3(nomDoc: str, placeDuFichier: str):
    for i in range(1, inputpage.max_row):
        if cell[i][11].value == "当天":
            if len(str(datetime.datetime.now().month)) == 1:
                datetimes = str(datetime.datetime.now().day) + "/0" + str(datetime.datetime.now().month) + "/" + str(
                    datetime.datetime.now().year)
            else:
                datetimes = str(datetime.datetime.now().day) + "/" + str(datetime.datetime.now().month) + "/" + str(
                    datetime.datetime.now().year)
        else:
            datetimes = cell[i][11].value

        chaine = cell[i][3].value + " sis " + cell[i][4].value + " " + cell[i][6].value + " " + cell[i][7].value + " " + \
                 str(cell[i][5].value) + " " + cell[i][8].value
        textlist = [cell[i][9].value, chaine, cell[i][10].value, datetimes, chaine, datetimes]

        doc = Document("授权书\\" + nomDoc)
        count = 0
        for p in doc.paragraphs:
            if '****' in p.text:
                inline = p.runs
                for j in range(len(inline)):
                    if '****' in inline[j].text:
                        text = inline[j].text.replace('****', textlist[count])
                        inline[j].text = text
                        count += 1
        doc.save(placeDuFichier + "\\" + str(cell[i][0].value) + "-Mandat-" + cell[i][3].value.upper() + ".docx")


def outputdocument4(nomDoc: str, placeDuFichier: str):
    for i in range(1, inputpage.max_row):

        textlist = [cell[i][5].value, cell[i][7].value, cell[i][8].value, cell[i][9].value, cell[i][2].value,
                    cell[i][3].value, cell[i][7].value, cell[i][3].value]

        doc = Document("授权书\\" + nomDoc)
        count = 0
        for p in doc.paragraphs:
            if '****' in p.text:
                inline = p.runs
                for j in range(len(inline)):
                    if '****' in inline[j].text:
                        text = inline[j].text.replace('****', textlist[count])
                        inline[j].text = text
                        count += 1
        doc.save(placeDuFichier  + str(cell[i][0].value) + "-Mandat-" + cell[i][7].value.upper() + ".docx")


def paris02f():
    nomDoc = "Paris-02法语授权书.docx"
    placeDuFichier = "输出\\Paris-02\\"
    outputdocument1(nomDoc, placeDuFichier)
    window.quit()


def paris16f():
    nomDoc = "Paris-16法语授权书.docx"
    placeDuFichier = "输出\\Paris-16\\"
    outputdocument2(nomDoc, placeDuFichier)
    window.quit()


def cscf():
    nomDoc = "Châtillon-sur-Cluses法语授权书.docx"
    placeDuFichier = "输出\\Châtillon-sur-Cluses\\"
    outputdocument3(nomDoc, placeDuFichier)
    window.quit()


def changeparis02f():
    nomDoc = "税务代表变更授权书 Paris 02.docx"
    placeDuFichier = "输出\\税务代表变更授权书 Paris 02\\"
    outputdocument4(nomDoc, placeDuFichier)
    window.quit()


path = os.getcwd()
fileList = os.listdir(path + "\模板")
fileName:str = None

n = 0
for i in fileList:
    if "模板" in fileList[n]:
        fileName = fileList[n]
        break
    n += 1
input = load_workbook("模板\\" + fileName)
inputpage = input.active
cell = tuple(inputpage)

window = tk.Tk()
window.title('授权书填写')
window.geometry('250x130')
paris02b = tk.Button(window, text='Paris-02-法语授权书', command=paris02f)
paris02b.pack()
paris16b = tk.Button(window, text='Paris-16-法语授权书', command=paris16f)
paris16b.pack()
cscb = tk.Button(window, text='Châtillon-sur-Cluses法语授权书', command=cscf)
cscb.pack()
changeparis02b = tk.Button(window, text='税务代表变更授权书 Paris 02', command=changeparis02f)
changeparis02b.pack()

window.mainloop()
